import datetime
import os
import smtplib
import requests
from bs4 import BeautifulSoup
import re
from docx import Document
from urllib import request as urlrequest

page = 1
u = []
link_array = []
parse_text = []
def input_link(page):
    link = 'http://www.banki.ru/news/lenta/page'+(str(page))+'/'
    return link
    
def get_url(link):
    proxy_host = '163.172.27.213:3128'
    req = urlrequest.Request(link)
    req.set_proxy(proxy_host, 'http')
    response = urlrequest.urlopen(req)
    link_1 = response.read().decode('utf8')
    # url = requests.get(link_1).text
    # print(link_1)
    return link_1
    
def get_soup(link_1):
    soup = BeautifulSoup(link_1, 'lxml')
    news = soup.find_all(class_='widget margin-bottom-default')
    # print(news)
    return news

# get_soup(get_url(input_link(page)))

def get_link_array(date):
    for i in date:
        x = datetime.datetime.strptime(i.get_text(), "%d.%m.%Y").date()
        ul = i.find_next('ul').find_all('li')
        for i in ul:
            href = i.find('a').get('href')
            time = i.find('span').text
            y = datetime.datetime.strptime(time, "%H:%M").time()
            u.append({'date': datetime.datetime.combine(x, y),
                    'url': href})
            # print(u)

def last_date_on_page(u):
    date_page = (list((u[-1]).values())[0]).date()
    return date_page
    
def delta_time():
    d1 = datetime.datetime.now()
    d2 = datetime.timedelta(days = 1)
    delta = (d1 - d2).date()
    return delta

def date_range(u):
    date_today = datetime.datetime.today()
    date_yearstaday = date_today - datetime.timedelta(days = 1)
    for i in u:
        if date_yearstaday < i['date'] <= date_today:
            if re.match(r'^/news', i['url']):
                if i['url']:
                    link_array.append("".join('http://www.banki.ru' + i['url']))
                    # return link_array

def array_list(link_array):
    for i in link_array:
        url = get_url(i)
        soup = BeautifulSoup(url, 'lxml')
        body = soup.find(class_='layout-column-center news-item')
        head = body.find(class_='header-h0').text.rstrip()
        date = body.find('time').text
        parse_text.append({
            'head': head,
            'date': date,
            'text': body.find(class_='article-text').text
            })
    parse_text.reverse()
    print(parse_text)

def to_word(parse_text):
    # filepath = '/Users/Korren/'
    y = datetime.datetime.today().strftime("%d-%m-%Y")
    document = Document()
    for i in parse_text:
        document.add_heading(i['head'], level=1)
        document.add_paragraph(i['date'])
        document.add_paragraph(i['text'])
    document.save(datetime.datetime.today().strftime("%d-%m-%Y") + '.docx')
    
# def send_email():
# from email.mime.text import MIMEText
# from email.mime.multipart import MIMEMultipart
# from email.mime.base import MIMEBase
# from email import encoders

# email_user = '***'
# email_send = '***'
# subject = 'python'

# msg = MIMEMultipart()
# msg['From'] = email_user
# msg['To'] = email_send
# msg['Subject'] = subject

# body = 'test'
# msg.attach(MIMEText(body, 'plain'))


# filename = (datetime.datetime.today().strftime("%d-%m-%Y") + '.docx')
# attachment = open(filename, 'rb')
# print(attachment)
# part = MIMEBase('application', 'octat-stream')
# part.set_payload(attachment)
# encoders.encode_base64(part)
# part.add_header('Content-Dispositions', "attachment; filename="+filename)

# msg.attach(part)
# text = msg.as_string()


# server = smtplib.SMTP('smtp.yandex.ru', 25)
# server.starttls()
# server.login(email_user, '***')

# server.sendmail(email_user, email_send)
# server.quit()

# #     # file_path = '/'
    


while not u or last_date_on_page(u) >= delta_time():
    for i in get_soup(get_url(input_link(page))):
        date_on_page = len(i.find_all('time'))
        date = i.find_all('time')
        date_today = datetime.datetime.now().date()
        first_date_on_page = datetime.datetime.strptime(i.find('time').get_text(), "%d.%m.%Y").date()
        if date_on_page > 2 and first_date_on_page == date_today:
            get_link_array(date[:2])
        else:
            get_link_array(date)
            page += 1
            
# print(u)
date_range(u)
# print(link_array)
array_list(link_array)
to_word(parse_text)
# send_email()


    